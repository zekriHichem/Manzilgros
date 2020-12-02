
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Cm
from docx.shared import Pt
import pandas as pd

def create_table(document):
    table = document.add_table(1,6)
    table.style = 'Table Grid'
    row = table.rows[0]
    cell = row.cells[0]
    table.columns[0].width = Cm(1.2)
    cell.text = "date"
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell = row.cells[1]
    table.columns[1].width = Cm(1)
    cell.text = "Référance"
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell = row.cells[2]
    table.columns[2].width = Cm(2)
    cell.text = "Désignation "
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell = row.cells[3]
    table.columns[3].width = Cm(1)
    cell.text = "Quantité"
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell = row.cells[4]
    table.columns[4].width = Cm(1)
    cell.text = "Prix vente TTC"
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell = row.cells[5]
    table.columns[5].width = Cm(1)
    cell.text = "total TTC"
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    return table


def create_row(table, product):
    row = table.add_row()
    cell = row.cells[0]
    cell.text = product.date
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell = row.cells[1]
    cell.text = product.ref
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell = row.cells[2]
    cell.text = product.designation
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell = row.cells[3]
    cell.text = str(product.quantite)
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell = row.cells[4]
    cell.text = str(product.prix_vent)
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell = row.cells[5]
    cell.text = str(product.prix_t_ttc)
    cell.paragraphs[0].runs[0].font.size = Pt(10)


def create_signature(document):
    signature = document.add_table(3, 1)
    row = signature.rows[0]
    cell = row.cells[0]
    cell.text = "Date: ............."
    cell.paragraphs[0].runs[0].font.size = Pt(14)

    row = signature.rows[1]
    cell = row.cells[0]
    cell.text = "Nom/Signature de "
    cell.paragraphs[0].runs[0].font.size = Pt(14)

    row = signature.rows[2]
    cell = row.cells[0]
    cell.text = "la structure expéditrice"
    cell.paragraphs[0].runs[0].font.size = Pt(14)


def run(products ):
    document = Document('../media/demo.docx')
    date = document.add_paragraph('')
    date.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date.add_run('Date: ..................').font.size = Pt(14)
    titre = document.add_paragraph('')
    titre.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    titre.add_run('BON DE TRANSFERT INTERNE N°: .....').font.size = Pt(24)
    ste = document.add_paragraph('')
    ste.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    ste_text = ste.add_run('STRUCTURE EXPDITRICE:')
    ste_text.font.underline = True
    ste_text.font.size = Pt(14)
    stre = document.add_paragraph('')
    stre.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    stre_text = stre.add_run('STRUCTURE RECEPTRICE:')
    stre_text.font.underline = True
    stre_text.font.size = Pt(14)
    table = create_table(document)
    i = 0
    k = 0
    for product in products:
        create_row(table, product)
        print(i)
        if i == 5:
            document.add_paragraph()
            create_signature(document)
            document.add_paragraph()
            document.add_paragraph()
            document.add_paragraph()
            table = create_table(document)
            k = 0



        i = i + 1
        k = k + 1
    document.save('facture.docx')

    return True

class Product():
    def __init__(self, date,codebar, ref, designation, quantite, prix_t_ttc, prix_vent):
        self.date = date
        self.codebar = codebar
        self.ref = ref
        self.designation = designation
        self.quantite = quantite
        self.prix_vent = prix_vent
        self.prix_t_ttc = prix_t_ttc

    @staticmethod
    def read_excel(url):
        datafram = pd.read_excel(url)
        products = []
        for d in datafram.values :
            product = Product(str(d[0]),str(d[1]) ,str(d[2]), str(d[3]), int(d[4]), float(d[6]), float(d[5]))
            products.append(product)
        return products


p = Product.read_excel('t.xlsx')

run(p)